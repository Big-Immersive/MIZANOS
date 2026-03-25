"""Generate downloadable .docx reports matching the daily briefing format."""

import io
import json
import logging
from datetime import datetime, timezone
from uuid import UUID

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy.ext.asyncio import AsyncSession

from apps.api.services.llm_config import get_llm_config
from apps.api.services.report_service import ReportService

logger = logging.getLogger(__name__)


class ReportDocumentService:
    """Generate a Word document report for selected projects."""

    def __init__(self, session: AsyncSession) -> None:
        self.session = session

    async def generate(self, product_ids: list[UUID]) -> io.BytesIO:
        """Build a .docx report for the given product IDs."""
        svc = ReportService(self.session)
        summary = await svc.get_summary()

        # Filter to selected projects
        projects = [p for p in summary["projects"] if p["product_id"] in product_ids]
        if not projects:
            projects = summary["projects"]

        ai_summary = await self._generate_executive_summary(projects)
        project_updates = await self._generate_project_updates(projects)

        doc = Document()
        self._set_default_font(doc)
        self._add_title(doc)
        self._add_executive_summary(doc, ai_summary)
        self._add_project_updates(doc, projects, project_updates)
        self._add_portfolio_table(doc, projects)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    # ------------------------------------------------------------------
    # Document sections
    # ------------------------------------------------------------------

    @staticmethod
    def _set_default_font(doc: Document) -> None:
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(10)

    @staticmethod
    def _add_title(doc: Document) -> None:
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = title.add_run("PROJECT STATUS UPDATE")
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0, 0, 0)

        date_str = datetime.now(timezone.utc).strftime("Daily Briefing — %d %B %Y")
        sub = doc.add_paragraph()
        run = sub.add_run(date_str)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(100, 100, 100)
        doc.add_paragraph()

    @staticmethod
    def _add_executive_summary(doc: Document, text: str) -> None:
        heading = doc.add_paragraph()
        run = heading.add_run("Executive Summary")
        run.bold = True
        run.font.size = Pt(14)

        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(10)
        doc.add_paragraph()

    @staticmethod
    def _add_project_updates(
        doc: Document, projects: list[dict], updates: dict[str, list[str]],
    ) -> None:
        heading = doc.add_paragraph()
        run = heading.add_run("Project Updates")
        run.bold = True
        run.font.size = Pt(14)

        for proj in projects:
            pid = str(proj["product_id"])
            name = proj["product_name"]
            stage = proj.get("stage") or "N/A"
            pm = proj.get("pm_name") or "—"

            # Project header: "Name  STATUS  PM: Name"
            ph = doc.add_paragraph()
            run = ph.add_run(f"{name}  ")
            run.bold = True
            run.font.size = Pt(12)

            run = ph.add_run(f"{stage.upper()}  ")
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 128, 0)

            run = ph.add_run(f"PM: {pm}")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(100, 100, 100)

            # Bullet points
            bullets = updates.get(pid, [])
            for bullet in bullets:
                bp = doc.add_paragraph(style="List Bullet")
                run = bp.add_run(bullet)
                run.font.size = Pt(10)

            doc.add_paragraph()

    @staticmethod
    def _add_portfolio_table(doc: Document, projects: list[dict]) -> None:
        heading = doc.add_paragraph()
        run = heading.add_run("Portfolio Directory")
        run.bold = True
        run.font.size = Pt(14)

        headers = ["Project", "Status", "PM", "Dev", "Tasks", "Commits", "Created"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"

        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(9)

        for proj in projects:
            row = table.add_row()
            cells = row.cells
            cells[0].text = proj["product_name"]
            cells[1].text = proj.get("stage") or "—"
            cells[2].text = proj.get("pm_name") or "—"
            cells[3].text = proj.get("dev_name") or "—"
            cells[4].text = f"{proj['completed_tasks']}/{proj['total_tasks']}"
            cells[5].text = str(proj.get("total_commits", 0))
            created = proj.get("created_at")
            if created:
                if isinstance(created, str):
                    created = datetime.fromisoformat(created)
                cells[6].text = created.strftime("%d %b %Y")
            else:
                cells[6].text = "—"

            for cell in cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)

    # ------------------------------------------------------------------
    # AI generation
    # ------------------------------------------------------------------

    async def _generate_executive_summary(self, projects: list[dict]) -> str:
        lines = []
        for p in projects:
            lines.append(
                f"- {p['product_name']} ({p.get('stage', 'N/A')}): "
                f"{p['completed_tasks']}/{p['total_tasks']} tasks done, "
                f"{p.get('total_commits', 0)} commits"
            )
        context = "\n".join(lines)

        prompt = (
            "Write a concise executive summary (1 paragraph, 4-6 sentences) for a "
            "daily project status briefing. Summarize the key highlights across these "
            "projects. Be specific about project names and achievements. "
            "No markdown, no bullet points — just flowing prose.\n\n"
            f"Projects:\n{context}"
        )
        return await self._call_llm(prompt)

    async def _generate_project_updates(
        self, projects: list[dict],
    ) -> dict[str, list[str]]:
        context = json.dumps(
            [
                {
                    "id": str(p["product_id"]),
                    "name": p["product_name"],
                    "stage": p.get("stage"),
                    "tasks_done": p["completed_tasks"],
                    "total_tasks": p["total_tasks"],
                    "in_progress": p["in_progress_tasks"],
                    "commits": p.get("total_commits", 0),
                }
                for p in projects
            ],
            indent=2,
        )

        prompt = (
            "For each project, generate 3-4 concise bullet points summarizing "
            "current status and recent progress. Return ONLY valid JSON with format:\n"
            '{"<product_id>": ["bullet 1", "bullet 2", ...], ...}\n'
            "No markdown fences. No explanation.\n\n"
            f"Projects:\n{context}"
        )
        raw = await self._call_llm(prompt)
        try:
            return json.loads(raw)
        except json.JSONDecodeError:
            logger.warning("Failed to parse project updates JSON")
            return {}

    async def _call_llm(self, prompt: str) -> str:
        import openai

        config = await get_llm_config(self.session)
        client = openai.AsyncOpenAI(api_key=config.api_key, base_url=config.base_url)
        response = await client.chat.completions.create(
            model=config.model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5,
            max_tokens=2048,
        )
        return response.choices[0].message.content or ""
